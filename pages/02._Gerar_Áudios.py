import streamlit as st
from docx import Document
import io
import os
import zipfile
import shutil
from elevenlabs import generate, Voice, VoiceSettings, set_api_key, RateLimitError
import docx

# Funções definidas anteriormente
def dividir_texto(texto, limite, separador='.'):
    partes = []
    parte_atual = ""

    for palavra in texto.split():
        if len(parte_atual + ' ' + palavra) <= limite:
            parte_atual += ' ' + palavra
        else:
            if parte_atual.endswith(separador):
                partes.append(parte_atual)
                parte_atual = palavra
            else:
                parte_atual += ' ' + palavra
    if parte_atual:
        partes.append(parte_atual)

    return partes

# Função principal do Streamlit
def streamlit_app():
    st.title("Gerador de Narração")

    # Campos para chave da API e voice_id
    api_key = st.text_input("Chave da API ElevenLabs", type="password")
    voice_id = st.text_input("Voice ID")

    # Upload do documento Word
    doc_file = st.file_uploader("Escolha o arquivo DOCX", type="docx")
    doc = docx.Document(doc_file)
    texto_arquivo = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    st.text_area("Texto do Arquivo", texto_arquivo, height=300)

    # Diretório temporário para os arquivos de áudio
    audio_dir = "audio_narracao_1"
    if not os.path.exists(audio_dir):
        os.makedirs(audio_dir)


    

    if doc_file is not None and api_key and voice_id:
        # Ler o arquivo DOCX da memória
        doc_memory = io.BytesIO(doc_file.getvalue())
        doc = Document(doc_memory)
        set_api_key(api_key)

        with st.expander("Visualizar arquivos criados"):
        # Botão para atualizar a lista de arquivos
            refresh_button = st.button("Atualizar Pasta")

            if refresh_button or 'refresh' not in st.session_state:
                st.session_state['refresh'] = True

            if st.session_state['refresh']:
                if os.path.exists(audio_dir):
                    for file in os.listdir(audio_dir):
                        audio_path = os.path.join(audio_dir, file)
                        with open(audio_path, "rb") as audio_file:
                            st.audio(audio_file.read(), format="audio/mp3")
                        if st.button(f"Excluir {file}"):
                            os.remove(audio_path)
                            st.session_state['refresh'] = False  # Requer atualização

        start_button = st.button("Iniciar Criação dos Áudios")

        col1, col2, col3 = st.columns(3)

        # Coloca o botão de download do ZIP na primeira coluna
        with col1:
            if st.button("📥 Baixar narrações em ZIP"):
                zip_memory = io.BytesIO()
                with zipfile.ZipFile(zip_memory, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for file in os.listdir(audio_dir):
                        zipf.write(os.path.join(audio_dir, file), file)
                zip_memory.seek(0)
                st.download_button(
                    label="Baixar ZIP",
                    data=zip_memory,
                    file_name="narrações.zip",
                    mime="application/zip"
                )

        # Coloca o botão para interromper a criação dos áudios na segunda coluna
        with col2:
            stop_button = st.button("🛑 Parar Criação dos Áudios")

            if stop_button:
                st.warning("A criação dos áudios foi interrompida.")
                return  # Interrompe a execução do script

        # Coloca o botão para limpar a pasta na terceira coluna
        with col3:
            if st.button("🧹 Limpar Pasta de Áudios e ZIP",type="primary"):
                shutil.rmtree(audio_dir)
                os.makedirs(audio_dir)
                st.success("Pasta de áudios e arquivo ZIP limpos com sucesso.")
        
        try:
            if start_button:
                # Processa o documento
                with st.spinner("Gerando áudios..."):
                    current_slide = 0
                    file_count = 0
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text.startswith("SLIDE:"):
                            current_slide = int(text.split(': ')[1])
                            file_count = 0
                        elif text and text != '.':
                            partes = dividir_texto(text, 200, '.')
                            for parte in partes:
                                file_count += 1
                                if parte.strip() and parte.strip() != '.':
                                    arquivo_audio = f"{audio_dir}/{current_slide}.{file_count}_narracao_slide.mp3"
                                    if not os.path.exists(arquivo_audio):
                                        audio = generate(
                                            text=parte,
                                            voice=Voice(
                                                voice_id=voice_id,
                                                settings=VoiceSettings(stability=1.0, similarity_boost=0.70, style=0.0, use_speaker_boost=True)
                                            ),
                                            model="eleven_multilingual_v2"
                                        )
                                        with open(arquivo_audio, 'wb') as file:
                                            file.write(audio)
                                        st.audio(arquivo_audio, format="audio/mp3")
                                        st.success(f"Texto criado para Slide {current_slide}, Parte {file_count}:\n\n{parte}")
            
                st.success("Todos os arquivos de áudio foram gerados com sucesso")

       
        except RateLimitError as e:
            st.error("Os créditos da API acabaram. É necessário trocar de conta ou aguardar a renovação dos créditos para continuar.")

    

if __name__ == "__main__":
    streamlit_app()
